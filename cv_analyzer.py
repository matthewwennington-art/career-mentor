"""
CV Analyzer Module
Compares CV with job listings and provides rating and improvement suggestions.
"""

import re
import os
from typing import Dict, List, Tuple
from collections import Counter
# #region agent log
try:
    import json
    from datetime import datetime
    with open(r'c:\Users\Matt\Google Drive\Stuff\Cursor Apps\Mentor\.cursor\debug.log', 'a', encoding='utf-8') as f:
        f.write(json.dumps({"sessionId":"debug-session","runId":"run1","hypothesisId":"E","location":"cv_analyzer.py:10","message":"Before importing PyPDF2 in cv_analyzer","data":{},"timestamp":int(datetime.now().timestamp()*1000)}) + '\n')
except: pass
# #endregion
try:
    # Try importing pypdf (new package name) first
    import pypdf as PyPDF2
    # #region agent log
    try:
        with open(r'c:\Users\Matt\Google Drive\Stuff\Cursor Apps\Mentor\.cursor\debug.log', 'a', encoding='utf-8') as f:
            f.write(json.dumps({"sessionId":"debug-session","runId":"run1","hypothesisId":"E","location":"cv_analyzer.py:13","message":"pypdf imported successfully as PyPDF2 in cv_analyzer","data":{"version":getattr(PyPDF2,'__version__','unknown')},"timestamp":int(datetime.now().timestamp()*1000)}) + '\n')
    except: pass
    # #endregion
except ImportError:
    # Fallback to PyPDF2 for older installations
    try:
        import PyPDF2
        # #region agent log
        try:
            with open(r'c:\Users\Matt\Google Drive\Stuff\Cursor Apps\Mentor\.cursor\debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({"sessionId":"debug-session","runId":"run1","hypothesisId":"E","location":"cv_analyzer.py:18","message":"PyPDF2 imported successfully (legacy) in cv_analyzer","data":{"version":getattr(PyPDF2,'__version__','unknown')},"timestamp":int(datetime.now().timestamp()*1000)}) + '\n')
        except: pass
        # #endregion
    except ImportError as e:
        # #region agent log
        try:
            with open(r'c:\Users\Matt\Google Drive\Stuff\Cursor Apps\Mentor\.cursor\debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({"sessionId":"debug-session","runId":"run1","hypothesisId":"E","location":"cv_analyzer.py:22","message":"Both pypdf and PyPDF2 import failed in cv_analyzer","data":{"error":str(e),"error_type":type(e).__name__},"timestamp":int(datetime.now().timestamp()*1000)}) + '\n')
        except: pass
        # #endregion
        raise ImportError(f"PDF library not found. Please install it with: pip install pypdf. Original error: {e}")
from docx import Document


class CVAnalyzer:
    """Analyzes CV against job listings and provides feedback."""
    
    def __init__(self):
        self.cv_text = ""
        self.job_listing_text = ""
        self.cv_keywords = []
        self.job_keywords = []
        self.match_score = 0.0
        self.suggestions = []
    
    def load_cv(self, filepath: str) -> bool:
        """Load CV from file (supports .txt, .pdf, .docx)."""
        try:
            if filepath.endswith('.pdf'):
                with open(filepath, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    self.cv_text = ""
                    for page in pdf_reader.pages:
                        self.cv_text += page.extract_text() + "\n"
            elif filepath.endswith('.docx'):
                doc = Document(filepath)
                self.cv_text = "\n".join([para.text for para in doc.paragraphs])
            elif filepath.endswith('.txt'):
                with open(filepath, 'r', encoding='utf-8') as f:
                    self.cv_text = f.read()
            else:
                print(f"Unsupported file format. Please use .txt, .pdf, or .docx")
                return False
            
            self.cv_text = self._clean_text(self.cv_text)
            self.cv_keywords = self._extract_keywords(self.cv_text)
            return True
        except Exception as e:
            print(f"Error loading CV: {str(e)}")
            return False
    
    def load_job_listing(self, filepath: str = None, text: str = None) -> bool:
        """Load job listing from file or text input."""
        try:
            if filepath:
                if filepath.endswith('.txt'):
                    with open(filepath, 'r', encoding='utf-8') as f:
                        self.job_listing_text = f.read()
                elif filepath.endswith('.pdf'):
                    with open(filepath, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        self.job_listing_text = ""
                        for page in pdf_reader.pages:
                            self.job_listing_text += page.extract_text() + "\n"
                elif filepath.endswith('.docx'):
                    doc = Document(filepath)
                    self.job_listing_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    print("Unsupported file format. Please use .txt, .pdf, or .docx")
                    return False
            elif text:
                self.job_listing_text = text
            else:
                print("Please provide either a filepath or text input")
                return False
            
            self.job_listing_text = self._clean_text(self.job_listing_text)
            self.job_keywords = self._extract_keywords(self.job_listing_text)
            return True
        except Exception as e:
            print(f"Error loading job listing: {str(e)}")
            return False
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Convert to lowercase for keyword matching
        return text.lower()
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract keywords from text (skills, technologies, qualifications)."""
        # Common technical skills and keywords
        common_skills = [
            'python', 'java', 'javascript', 'sql', 'html', 'css', 'react', 'angular',
            'node.js', 'django', 'flask', 'aws', 'azure', 'docker', 'kubernetes',
            'git', 'agile', 'scrum', 'project management', 'leadership', 'communication',
            'analytics', 'data analysis', 'machine learning', 'ai', 'cloud computing',
            'devops', 'ci/cd', 'rest api', 'microservices', 'database', 'nosql',
            'excel', 'powerpoint', 'presentation', 'negotiation', 'sales', 'marketing',
            'finance', 'accounting', 'design', 'ui/ux', 'customer service', 'teamwork'
        ]
        
        # Extract words that appear in common skills
        keywords = []
        text_lower = text.lower()
        for skill in common_skills:
            if skill in text_lower:
                keywords.append(skill)
        
        # Also extract capitalized terms (likely proper nouns, technologies, etc.)
        capitalized_terms = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        keywords.extend([term.lower() for term in capitalized_terms if len(term) > 3])
        
        # Extract years of experience patterns
        experience_patterns = re.findall(r'\d+\+?\s*years?\s*(?:of\s*)?experience', text_lower)
        keywords.extend(experience_patterns)
        
        return list(set(keywords))  # Remove duplicates
    
    def analyze_match(self) -> Dict:
        """Analyze how well CV matches the job listing."""
        if not self.cv_text or not self.job_listing_text:
            return {"error": "Please load both CV and job listing first"}
        
        # Calculate keyword match
        cv_keyword_set = set(self.cv_keywords)
        job_keyword_set = set(self.job_keywords)
        
        matching_keywords = cv_keyword_set.intersection(job_keyword_set)
        missing_keywords = job_keyword_set - cv_keyword_set
        
        # Calculate match score (0-100)
        if len(job_keyword_set) > 0:
            self.match_score = (len(matching_keywords) / len(job_keyword_set)) * 100
        else:
            self.match_score = 0
        
        # Generate suggestions
        self.suggestions = self._generate_suggestions(matching_keywords, missing_keywords)
        
        return {
            "match_score": round(self.match_score, 2),
            "matching_keywords": list(matching_keywords),
            "missing_keywords": list(missing_keywords),
            "suggestions": self.suggestions
        }
    
    def _generate_suggestions(self, matching: set, missing: set) -> List[str]:
        """Generate improvement suggestions."""
        suggestions = []
        
        # Score-based suggestions
        if self.match_score < 30:
            suggestions.append("Your CV has a low match score. Consider highlighting more relevant skills and experiences.")
        elif self.match_score < 50:
            suggestions.append("Your CV has a moderate match. There's room for improvement to better align with the job requirements.")
        elif self.match_score < 70:
            suggestions.append("Your CV has a good match. A few enhancements could make it even stronger.")
        else:
            suggestions.append("Your CV has a strong match with the job listing. Great alignment!")
        
        # Missing keywords suggestions
        if missing:
            top_missing = list(missing)[:5]  # Top 5 missing keywords
            suggestions.append(f"Consider adding or emphasizing these keywords from the job listing: {', '.join(top_missing)}")
        
        # Structure suggestions
        suggestions.append("Ensure your CV clearly highlights your most relevant experiences at the top.")
        suggestions.append("Use action verbs and quantify achievements where possible (e.g., 'Increased sales by 25%').")
        suggestions.append("Tailor your CV summary/objective to match the job description's key requirements.")
        
        # Format suggestions
        suggestions.append("Make sure your CV is well-formatted, easy to read, and free of typos.")
        suggestions.append("Include a skills section that matches the job requirements.")
        
        return suggestions
    
    def display_analysis(self):
        """Display the analysis results."""
        analysis = self.analyze_match()
        
        if "error" in analysis:
            print(analysis["error"])
            return
        
        print("\n" + "="*60)
        print("CV ANALYSIS RESULTS")
        print("="*60)
        
        print(f"\nMatch Score: {analysis['match_score']:.1f}/100")
        
        if analysis['match_score'] >= 70:
            print("Status: ✅ Strong Match")
        elif analysis['match_score'] >= 50:
            print("Status: ⚠️  Moderate Match")
        else:
            print("Status: ❌ Needs Improvement")
        
        print(f"\nMatching Keywords ({len(analysis['matching_keywords'])}):")
        if analysis['matching_keywords']:
            for keyword in analysis['matching_keywords'][:10]:  # Show top 10
                print(f"  ✓ {keyword}")
        else:
            print("  None found")
        
        print(f"\nMissing Keywords ({len(analysis['missing_keywords'])}):")
        if analysis['missing_keywords']:
            for keyword in list(analysis['missing_keywords'])[:10]:  # Show top 10
                print(f"  ✗ {keyword}")
        else:
            print("  None - great alignment!")
        
        print("\n" + "="*60)
        print("IMPROVEMENT SUGGESTIONS")
        print("="*60)
        for i, suggestion in enumerate(analysis['suggestions'], 1):
            print(f"{i}. {suggestion}")
        
        print("\n" + "="*60)

